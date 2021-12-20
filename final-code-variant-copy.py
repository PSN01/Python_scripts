#!/usr/bin/env python
# coding: utf-8

# In[3]:


'''code for copying results in a dataframe to a table column in a document'''
import pandas as pd 
import docx
doc = docx.Document('Variant_analysis_template.docx')
doc1=docx.Document('BiyyalaManjuja.docx')
table=doc.tables[1]
table1=doc1.tables[1]
data = [[cell.text for cell in row.cells] for row in table.rows]
df=pd.read_csv('predictions.csv', header=None)
df1 = pd.DataFrame(data)
df1.iloc[:,4]=df.iloc[:,0]
i=1
for i, column in enumerate(df1) :
        for row in range(df1.shape[0]) :
            table.cell(row, i).text = str(df1[column][row])
doc.save('Variant_analysis_template.docx')  


# In[ ]:




