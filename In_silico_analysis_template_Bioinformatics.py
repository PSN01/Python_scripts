#!/usr/bin/env python
# coding: utf-8

# In[35]:


'''code for copying results in a dataframe to a table column in a document'''
import pandas as pd 
import docx
import numpy as np
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor
doc = docx.Document('template.docx')
doc1=docx.Document('predictions.docx')
table=doc.tables[1]
table1=doc1.tables[0]
data = [[cell.text for cell in row.cells] for row in table.rows]
data1= [[cell.text for cell in row.cells] for row in table1.rows]
df1 = pd.DataFrame(data)
df2= pd.DataFrame(data1)
for column in df2.columns:
        df1.iloc[:,4]=df2.iloc[:,column]
for i, column in enumerate(df1):
        for row in range(df1.shape[0]):
            table.cell(row, i).text = str(df1[column][row])
for i in df1.columns:
    hdr_cells = table.rows[0].cells
    run = table.cell(0, i).paragraphs[0].runs[0]
    run.font.bold = True
for i in range(1,12):
    col_cells=table.rows[i].cells
    run=table.cell(i,4).paragraphs[0].runs[0]
    run.font.bold=True
    font = run.font
    font.color.rgb = RGBColor(255, 0, 0)
s=table.cell(0,4).text
print(s)
font = doc.styles['Normal'].font
font.name = 'Times New Roman'
font.size= Pt(9.5)
doc.save(s+'template.docx')


# In[ ]:




